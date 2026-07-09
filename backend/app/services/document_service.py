import io
import logging
import os
import tempfile
from typing import Optional
from uuid import uuid4

import pdfplumber
from fastapi import UploadFile
from openai import AsyncOpenAI
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import get_settings
from app.models.models import Document, DocumentChunk

settings = get_settings()

logger = logging.getLogger(__name__)

client = AsyncOpenAI(api_key=settings.OPENAI_API_KEY)


class DocumentService:
    def __init__(self, db: AsyncSession):
        self.db = db

    async def upload_and_process(
        self, chatbot_id: str, file: UploadFile, file_type: str
    ) -> Document:
        file_content = await file.read()
        file_size = len(file_content)
        filename = f"{uuid4().hex}.{file_type}"

        doc = Document(
            chatbot_id=chatbot_id,
            filename=filename,
            original_filename=file.filename or "unknown",
            file_type=file_type,
            file_size=file_size,
            status="processing",
        )
        self.db.add(doc)
        await self.db.flush()

        try:
            text = self._extract_text(file_content, file_type)
            chunks = self._split_text(text, max_chunk_size=1000, overlap=200)

            embeddings = await self._get_embeddings_batch(chunks)

            if settings.PINECONE_API_KEY:
                await self._store_pinecone(chatbot_id, doc.id, chunks, embeddings)
            else:
                await self._store_pgvector(doc.id, chunks, embeddings)

            doc.chunk_count = len(chunks)
            doc.status = "processed"
        except Exception as e:
            logger.error(f"Error processing document {filename}: {e}")
            doc.status = "error"
            doc.error_message = str(e)

        await self.db.flush()
        return doc

    def _extract_text(self, file_content: bytes, file_type: str) -> str:
        if file_type == "pdf":
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                text = ""
                for page in pdf.pages:
                    text += page.extract_text() or ""
                return text
        elif file_type in ("txt", "md", "csv"):
            return file_content.decode("utf-8", errors="ignore")
        elif file_type == "docx":
            from docx import Document as DocxDocument

            doc = DocxDocument(io.BytesIO(file_content))
            return "\n".join([p.text for p in doc.paragraphs])
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _split_text(
        self, text: str, max_chunk_size: int = 1000, overlap: int = 200
    ) -> list[str]:
        import tiktoken

        try:
            encoding = tiktoken.encoding_for_model("gpt-4o-mini")
        except KeyError:
            encoding = tiktoken.get_encoding("cl100k_base")

        tokens = encoding.encode(text)
        chunks = []
        start = 0

        while start < len(tokens):
            end = start + max_chunk_size
            chunk_tokens = tokens[start:end]
            chunk_text = encoding.decode(chunk_tokens)
            chunks.append(chunk_text)
            start = end - overlap

        return chunks

    async def _get_embeddings_batch(self, texts: list[str]) -> list[list[float]]:
        batch_size = 100
        all_embeddings = []

        for i in range(0, len(texts), batch_size):
            batch = texts[i : i + batch_size]
            response = await client.embeddings.create(
                model=settings.OPENAI_EMBEDDING_MODEL,
                input=batch,
            )
            all_embeddings.extend([item.embedding for item in response.data])

        return all_embeddings

    async def _store_pinecone(
        self,
        chatbot_id: str,
        document_id: str,
        chunks: list[str],
        embeddings: list[list[float]],
    ):
        from pinecone import Pinecone

        pc = Pinecone(api_key=settings.PINECONE_API_KEY)
        index = pc.Index(settings.PINECONE_INDEX_NAME)

        vectors = []
        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            chunk_id = f"{document_id}_{i}"
            vectors.append({
                "id": chunk_id,
                "values": embedding,
                "metadata": {
                    "content": chunk,
                    "document_id": str(document_id),
                    "chunk_index": i,
                },
            })

        index.upsert(vectors=vectors, namespace=str(chatbot_id))

        for i, chunk in enumerate(chunks):
            db_chunk = DocumentChunk(
                document_id=document_id,
                content=chunk,
                chunk_index=i,
                token_count=len(chunk.split()),
                pinecone_id=f"{document_id}_{i}",
            )
            self.db.add(db_chunk)

    async def _store_pgvector(
        self, document_id: str, chunks: list[str], embeddings: list[list[float]]
    ):
        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            db_chunk = DocumentChunk(
                document_id=document_id,
                content=chunk,
                chunk_index=i,
                token_count=len(chunk.split()),
            )
            self.db.add(db_chunk)

    async def delete_document(self, document_id: str, chatbot_id: str):
        result = await self.db.execute(
            select(Document).where(Document.id == document_id)
        )
        doc = result.scalar_one_or_none()
        if doc:
            if doc.pinecone_id and settings.PINECONE_API_KEY:
                from pinecone import Pinecone

                pc = Pinecone(api_key=settings.PINECONE_API_KEY)
                index = pc.Index(settings.PINECONE_INDEX_NAME)
                chunk_ids = [
                    f"{document_id}_{i}" for i in range(doc.chunk_count)
                ]
                index.delete(ids=chunk_ids, namespace=str(chatbot_id))

            await self.db.delete(doc)
            await self.db.flush()